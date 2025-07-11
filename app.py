import json
import os
import re
import tempfile
import time

import openai
import requests
import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document


def is_url_or_link(text):
    """Check if text is a URL or link that should not be translated."""
    # Common URL patterns
    url_patterns = [
        r"https?://[^\s]+",  # HTTP/HTTPS URLs
        r"www\.[^\s]+",  # www URLs
        r"[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",  # Domain names
        r"ftp://[^\s]+",  # FTP URLs
        r"mailto:[^\s]+",  # Email links
        r"file://[^\s]+",  # File URLs
    ]

    # Check if text matches any URL pattern
    for pattern in url_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    # Check for common link indicators
    link_indicators = ["http://", "https://", "www.", "mailto:", "ftp://", "file://"]
    text_lower = text.lower()
    for indicator in link_indicators:
        if indicator in text_lower:
            return True

    return False


def translate_with_openai_direct(text, src_lang, dest_lang, api_key):
    """Translate text using direct HTTP requests to OpenAI API."""
    try:
        # Clean the API key
        api_key = api_key.strip()

        # Validate API key format
        if not api_key.startswith("sk-"):
            raise Exception("Invalid API key format")

        # Create language mapping for better prompts
        lang_names = {
            "en": "English",
            "de": "German",
            "es": "Spanish",
            "fr": "French",
            "it": "Italian",
            "pt": "Portuguese",
            "ru": "Russian",
            "ja": "Japanese",
            "ko": "Korean",
            "zh": "Chinese",
            "ar": "Arabic",
            "hi": "Hindi",
        }

        src_lang_name = lang_names.get(src_lang, src_lang)
        dest_lang_name = lang_names.get(dest_lang, dest_lang)

        # Limit text length to avoid token limits
        if len(text) > 3000:
            text = text[:3000] + "..."

        # Prepare the request
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        data = {
            "model": "gpt-3.5-turbo",
            "messages": [
                {
                    "role": "system",
                    "content": f"You are a professional translator. Translate the following text from {src_lang_name} to {dest_lang_name}. Preserve the original formatting, punctuation, and meaning. Do not translate URLs, email addresses, or technical terms that should remain unchanged. Return only the translated text without any explanations.",
                },
                {"role": "user", "content": text},
            ],
            "max_tokens": 1000,
            "temperature": 0.3,
        }

        # Make the request
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30,
        )

        if response.status_code == 200:
            result = response.json()
            translated_text = result["choices"][0]["message"]["content"].strip()
            return translated_text if translated_text else text
        elif response.status_code == 401:
            raise Exception("Invalid API key or authentication failed")
        elif response.status_code == 429:
            raise Exception("Rate limit exceeded. Please wait a moment and try again.")
        else:
            raise Exception(
                f"API request failed with status {response.status_code}: {response.text}"
            )

    except requests.exceptions.RequestException as e:
        raise Exception(f"Network error: {str(e)}")
    except Exception as e:
        raise Exception(f"Direct API translation error: {str(e)}")


def translate_with_openai(text, src_lang, dest_lang, api_key):
    """Translate text using OpenAI API with fallback to direct HTTP."""
    try:
        # Clean the API key and create client
        api_key = api_key.strip()

        # Validate API key format
        if not api_key.startswith("sk-"):
            raise Exception("Invalid API key format")

        # Try using the OpenAI client first
        try:
            client = openai.OpenAI(api_key=api_key)

            # Create language mapping for better prompts
            lang_names = {
                "en": "English",
                "de": "German",
                "es": "Spanish",
                "fr": "French",
                "it": "Italian",
                "pt": "Portuguese",
                "ru": "Russian",
                "ja": "Japanese",
                "ko": "Korean",
                "zh": "Chinese",
                "ar": "Arabic",
                "hi": "Hindi",
            }

            src_lang_name = lang_names.get(src_lang, src_lang)
            dest_lang_name = lang_names.get(dest_lang, dest_lang)

            # Limit text length to avoid token limits
            if len(text) > 3000:
                text = text[:3000] + "..."

            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the following text from {src_lang_name} to {dest_lang_name}. Preserve the original formatting, punctuation, and meaning. Do not translate URLs, email addresses, or technical terms that should remain unchanged. Return only the translated text without any explanations.",
                    },
                    {"role": "user", "content": text},
                ],
                max_tokens=1000,
                temperature=0.3,
            )

            result = response.choices[0].message.content.strip()
            return result if result else text

        except (TypeError, AttributeError) as e:
            if "proxies" in str(e) or "unexpected keyword argument" in str(e):
                # Fallback to direct HTTP request
                return translate_with_openai_direct(text, src_lang, dest_lang, api_key)
            else:
                raise e

    except openai.AuthenticationError:
        raise Exception("Invalid API key or authentication failed")
    except openai.RateLimitError:
        raise Exception("Rate limit exceeded. Please wait a moment and try again.")
    except openai.APIError as e:
        raise Exception(f"OpenAI API error: {str(e)}")
    except Exception as e:
        raise Exception(f"OpenAI translation error: {str(e)}")


def translate_with_google(text, src_lang, dest_lang):
    """Translate text using Google Translate."""
    try:
        translator = GoogleTranslator(source=src_lang, target=dest_lang)
        result = translator.translate(text)
        return result
    except Exception as e:
        raise Exception(f"Google translation error: {str(e)}")


def translate_text_safely(
    text, src_lang, dest_lang, translation_method, openai_api_key=None, max_retries=3
):
    """Safely translate text with retry logic and error handling."""
    # Skip translation if text is a URL or link
    if is_url_or_link(text):
        return text

    # Skip translation if text is too short or empty
    if len(text.strip()) < 2:
        return text

    for attempt in range(max_retries):
        try:
            if translation_method == "OpenAI" and openai_api_key:
                return translate_with_openai(text, src_lang, dest_lang, openai_api_key)
            else:
                return translate_with_google(text, src_lang, dest_lang)
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(1)  # Wait before retrying
                continue
            else:
                st.warning(f"Failed to translate: '{text[:50]}...' - {str(e)}")
                return text  # Return original text if translation fails
    return text


def translate_docx(file, src_lang, dest_lang, translation_method, openai_api_key=None):
    doc = Document(file)

    # Progress tracking
    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    total_cells = sum(
        len([c for c in row.cells if c.text.strip()])
        for table in doc.tables
        for row in table.rows
    )
    total_items = total_paragraphs + total_cells

    progress_bar = st.progress(0)
    status_text = st.empty()

    current_item = 0

    # Translate paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            try:
                status_text.text(
                    f"Translating paragraph {current_item + 1}/{total_items}..."
                )
                translated_text = translate_text_safely(
                    para.text, src_lang, dest_lang, translation_method, openai_api_key
                )
                para.text = translated_text
                current_item += 1
                progress_bar.progress(current_item / total_items)
            except Exception as e:
                st.error(f"Error translating paragraph: {str(e)}")
                current_item += 1
                progress_bar.progress(current_item / total_items)
                continue

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    try:
                        status_text.text(
                            f"Translating table cell {current_item + 1}/{total_items}..."
                        )
                        translated_text = translate_text_safely(
                            cell.text,
                            src_lang,
                            dest_lang,
                            translation_method,
                            openai_api_key,
                        )
                        cell.text = translated_text
                        current_item += 1
                        progress_bar.progress(current_item / total_items)
                    except Exception as e:
                        st.error(f"Error translating table cell: {str(e)}")
                        current_item += 1
                        progress_bar.progress(current_item / total_items)
                        continue

    progress_bar.empty()
    status_text.empty()
    return doc


st.title("📄 DOCX Translator Tool")

# Sidebar for settings
with st.sidebar:
    st.header("⚙️ Translation Settings")

    translation_method = st.selectbox(
        "Translation Service",
        ["Google Translate", "OpenAI"],
        help="Choose your preferred translation service",
    )

    if translation_method == "OpenAI":
        st.info("🔑 OpenAI requires an API key for translation")
        # Try to get API key from secrets, fallback to text input
        try:
            openai_api_key = st.secrets.get("OPENAI_API_KEY")
        except:
            openai_api_key = None

        openai_api_key = openai_api_key or st.text_input(
            "OpenAI API Key",
            value="REMOVED_KEY",
            type="password",
            help="Enter your OpenAI API key. Get one at https://platform.openai.com/api-keys",
        )

        if not openai_api_key:
            st.warning("⚠️ Please enter your OpenAI API key to use OpenAI translation")
    else:
        openai_api_key = None

# Language codes reference
with st.expander("ℹ️ Language Codes Reference"):
    st.markdown(
        """
    **Common Language Codes:**
    - `en` - English
    - `de` - German  
    - `es` - Spanish
    - `fr` - French
    - `it` - Italian
    - `pt` - Portuguese
    - `ru` - Russian
    - `ja` - Japanese
    - `ko` - Korean
    - `zh` - Chinese
    - `ar` - Arabic
    - `hi` - Hindi
    
    Use 2-letter ISO codes, not full language names!
    
    **Note:** URLs and links will be preserved and not translated.
    """
    )

uploaded_file = st.file_uploader("Upload your .docx file", type=["docx"])
src_lang = st.text_input(
    "Source language code",
    placeholder="e.g., en, de, es",
    help="Enter the 2-letter language code of your source document",
)
dest_lang = st.text_input(
    "Destination language code",
    placeholder="e.g., es, fr, de",
    help="Enter the 2-letter language code for translation",
)

if uploaded_file and src_lang and dest_lang:
    # Check if OpenAI is selected but no API key provided
    if translation_method == "OpenAI" and not openai_api_key:
        st.error(
            "❌ Please provide your OpenAI API key in the sidebar to use OpenAI translation."
        )
    else:
        if st.button("Translate"):
            try:
                with st.spinner("Preparing translation..."):
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".docx"
                    ) as tmp_input:
                        tmp_input.write(uploaded_file.read())
                        tmp_input_path = tmp_input.name

                    doc = translate_docx(
                        tmp_input_path,
                        src_lang,
                        dest_lang,
                        translation_method,
                        openai_api_key,
                    )

                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=".docx"
                    ) as tmp_output:
                        doc.save(tmp_output.name)
                        tmp_output_path = tmp_output.name

                    with open(tmp_output_path, "rb") as f:
                        st.success("✅ Translation complete!")
                        st.download_button(
                            label="Download Translated DOCX",
                            data=f,
                            file_name=f"translated_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
            except Exception as e:
                st.error(f"An error occurred during translation: {str(e)}")
                st.info("Please try again or check your language codes.")
